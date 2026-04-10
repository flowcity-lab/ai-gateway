"""
import_parser.py — Intelligenter Import: Datei-Parser, Telefon-Normalisierung und KI-Spaltenanalyse

Endpoints:
  POST /import/parse             — CSV / XLSX / XLS / ODS parsen
  POST /import/normalize-phones  — Telefonnummern normalisieren + Länder erkennen
  POST /import/analyze-columns   — GPT-4o analysiert Spalten anhand von Werten + schlägt Custom Fields vor
"""

import asyncio
import io
import csv
import json
import base64
import logging
import os
import re
import time
import uuid
from typing import Optional

import chardet
import httpx
from fastapi import APIRouter, BackgroundTasks, HTTPException, Request
from pydantic import BaseModel
from openai import OpenAI as _OpenAI

log = logging.getLogger("ai-gateway.import")
router = APIRouter()

# Eigener OpenAI-Client (vermeidet circular import aus main.py)
_oai = _OpenAI(api_key=os.environ.get("OPENAI_API_KEY", ""))
_ANALYZE_MODEL = os.environ.get("IMPORT_ANALYZE_MODEL", "gpt-4o-mini")  # mini reicht hier vollständig

# Gateway-Secret für Callbacks (gleiche Variable wie main.py)
_GATEWAY_SECRET = os.environ.get("AI_GATEWAY_SECRET", "")

# In-Memory Task-Store für async Extraktion (TTL: 2h, Cleanup bei GET)
# { task_id: {"status": "processing"|"complete"|"failed", "result": dict|None, "error": str|None, "created_at": float} }
_import_tasks: dict = {}

# ── Pydantic Models ───────────────────────────────────────────────────────────

class ParseRequest(BaseModel):
    file_content_base64: str
    file_name: str
    file_type: Optional[str] = None

class NormalizePhoneRequest(BaseModel):
    phones: list[str]
    fallback_country: Optional[str] = None
    candidate_countries: list[str] = []

class AnalyzeColumnsRequest(BaseModel):
    headers: list[str]
    sample_rows: list[list[str]]
    existing_custom_field_names: list[str] = []

class ExtractContactsRequest(BaseModel):
    file_content_base64: str
    file_name: str
    file_type: str                              # csv, xlsx, xls, ods, pdf, jpg, png, webp, heic, tif, bmp, docx, rtf, odt, vcf, eml
    exclude_identity: dict = {}                 # Trainer-eigene Daten: names, emails, phones, company_name
    available_fields: list[str] = []            # System + Custom Fields die der Trainer hat
    chunk_index: int = 0                        # Intern für chunked processing
    model: str = "gpt-4.1"                     # Text-Extraktion: konfigurierbar über AI-Config
    vision_model: str = "gpt-4.1"             # Vision/Bild-Extraktion: konfigurierbar über AI-Config

class ExtractContactsAsyncRequest(ExtractContactsRequest):
    """Async-Variante: Gateway verarbeitet im Hintergrund, Ergebnis via Callback."""
    job_id: int                                 # AppJob.id — für Callback-Zuordnung
    callback_url: str = ""                      # Laravel: POST /api/ai/import/complete
    laravel_base_url: str = ""                  # Fallback: z.B. https://trainer.example.com

# ── Bekannte Export-Formate (Auto-Erkennung via Header-Patterns) ─────────────

KNOWN_FORMATS = {
    "mailchimp": ["email address", "member_rating", "optin_time"],
    "brevo":     ["email", "double opt-in"],
    "hubspot":   ["contact id", "hs_object_id", "createdate"],
}

EU_CANDIDATE_COUNTRIES = [
    "AT", "DE", "CH", "FR", "IT", "ES", "NL", "BE", "PL", "CZ",
    "SK", "HU", "RO", "BG", "HR", "SI", "LU", "PT", "SE", "DK",
    "FI", "NO", "GR", "IE", "LT", "LV", "EE",
]

# ── Hilfsfunktionen ───────────────────────────────────────────────────────────

def detect_format(headers: list[str]) -> str:
    """Erkennt bekannte Export-Formate anhand der Spaltenbezeichnungen."""
    lower_headers = {h.lower().strip() for h in headers}
    for fmt_name, signature in KNOWN_FORMATS.items():
        if all(sig in lower_headers for sig in signature):
            return fmt_name
    return "generic"


def parse_csv_bytes(raw_bytes: bytes) -> dict:
    """CSV parsen: Encoding und Delimiter auto-erkennen."""
    # Encoding erkennen
    detected = chardet.detect(raw_bytes)
    encoding = detected.get("encoding") or "utf-8"
    if encoding.lower() in ("ascii", "iso-8859-1", "latin-1", "windows-1252"):
        encoding = "utf-8-sig"  # BOM-tolerant

    try:
        text = raw_bytes.decode(encoding, errors="replace")
    except Exception:
        text = raw_bytes.decode("utf-8", errors="replace")

    # BOM entfernen
    text = text.lstrip("\ufeff")

    # Delimiter erkennen: Sniffer + Fallback-Liste
    delimiter = ","
    try:
        sample = text[:4096]
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        delimiter = dialect.delimiter
    except csv.Error:
        # Fallback: zählen welcher Delimiter am häufigsten vorkommt
        first_line = text.split("\n")[0] if "\n" in text else text[:200]
        counts = {d: first_line.count(d) for d in [",", ";", "\t", "|"]}
        delimiter = max(counts, key=counts.get)

    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows_raw = list(reader)

    # Leere Zeilen entfernen
    rows_raw = [r for r in rows_raw if any(cell.strip() for cell in r)]

    if not rows_raw:
        return {"headers": [], "rows": [], "row_count": 0, "has_data": False}

    headers = [h.strip() for h in rows_raw[0]]
    data_rows = [[cell.strip() for cell in row] for row in rows_raw[1:]]

    return {
        "name": "CSV",
        "headers": headers,
        "rows": data_rows,
        "row_count": len(data_rows),
        "has_data": len(data_rows) > 0,
        "detected_format": detect_format(headers),
        "delimiter_detected": delimiter,
        "encoding_detected": encoding,
    }


def parse_excel_bytes(raw_bytes: bytes) -> list[dict]:
    """Excel (XLSX/XLS/ODS) parsen: alle Sheets extrahieren."""
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(raw_bytes), read_only=True, data_only=True)
    sheets = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        all_rows = []
        for row in ws.iter_rows(values_only=True):
            # None-Werte zu leeren Strings, alles zu String konvertieren
            cleaned = [str(cell).strip() if cell is not None else "" for cell in row]
            if any(cell for cell in cleaned):  # nur nicht-leere Zeilen
                all_rows.append(cleaned)

        if not all_rows:
            sheets.append({
                "name": sheet_name,
                "headers": [],
                "rows": [],
                "row_count": 0,
                "has_data": False,
            })
            continue

        headers = all_rows[0]
        data_rows = all_rows[1:]

        sheets.append({
            "name": sheet_name,
            "headers": headers,
            "rows": data_rows,
            "row_count": len(data_rows),
            "has_data": len(data_rows) > 0,
            "detected_format": detect_format(headers),
        })

    wb.close()
    return sheets


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.post("/import/parse")
async def parse_import_file(payload: ParseRequest, request: Request):
    """
    Parst eine hochgeladene Datei (CSV oder Excel) und gibt strukturierte
    Sheet/Zeilen-Daten zurück. Kein LLM beteiligt — rein deterministisch.
    """
    try:
        raw_bytes = base64.b64decode(payload.file_content_base64)
    except Exception:
        raise HTTPException(400, "Ungültiger Base64-Inhalt")

    # Dateitype ermitteln
    file_type = (payload.file_type or "").lower()
    if not file_type:
        ext = payload.file_name.rsplit(".", 1)[-1].lower() if "." in payload.file_name else ""
        file_type = ext

    try:
        if file_type == "csv" or file_type in ("txt", "tsv"):
            sheet = parse_csv_bytes(raw_bytes)
            return {
                "file_type": "csv",
                "sheets": [sheet],
                "delimiter_detected": sheet.get("delimiter_detected"),
                "encoding_detected": sheet.get("encoding_detected"),
            }

        elif file_type in ("xlsx", "xls", "ods", "xlsm"):
            sheets = parse_excel_bytes(raw_bytes)
            return {
                "file_type": file_type,
                "sheets": sheets,
            }

        else:
            raise HTTPException(415, f"Nicht unterstütztes Dateiformat: {file_type}. Unterstützt: csv, xlsx, xls, ods")

    except HTTPException:
        raise
    except Exception as e:
        log.error("Import parse error: %s", str(e), exc_info=True)
        raise HTTPException(500, f"Fehler beim Parsen der Datei: {str(e)}")


@router.post("/import/normalize-phones")
async def normalize_phones(payload: NormalizePhoneRequest, request: Request):
    """
    Normalisiert Telefonnummern nach E.164 und erkennt das Ursprungsland.

    Rückgabe pro Nummer:
      - e164: "+436641234567" oder null wenn nicht normalisierbar
      - country: "AT" oder null
      - confidence: "high" (1 Land), "medium" (2-3 Länder), "low" (viele/keine)
      - possible_countries: alle möglichen Länder
      - national_format: lesbares Format
    """
    import phonenumbers

    candidates = payload.candidate_countries or EU_CANDIDATE_COUNTRIES
    fallback = (payload.fallback_country or "").upper() or None

    results = []
    for raw in payload.phones:
        result = _normalize_single(raw, fallback, candidates, phonenumbers)
        results.append({"raw": raw, **result})

    return {"results": results}


_ANALYZE_SYSTEM_PROMPT = """
Du bist ein Experte für Datenimport von Kontaktlisten. Deine Aufgabe: Spalten aus einer Excel- oder CSV-Datei
eines Trainers oder Coaches den richtigen Kontaktfeldern zuordnen.

Verfügbare System-Felder:
  first_name        — Vorname
  last_name         — Nachname
  full_name         — Vollständiger Name (wird automatisch in Vor- + Nachname gesplittet)
  email             — E-Mail-Adresse
  phone             — Telefonnummer (Mobil oder Festnetz)
  date_of_birth     — Geburtsdatum
  organization_name — Firmenname / Organisation / Arbeitgeber
  organization_role — Position / Funktion / Berufsbezeichnung in der Firma
  notes             — Notizen, Freitext, sonstige Bemerkungen
  doi_confirmed_at  — DOI-Bestätigungsdatum (Double Opt-In / Einwilligung)
  status            — Kontaktstatus: lead, interessent, kunde, inaktiv
  _ignore           — Spalte NICHT importieren

Regeln:
1. Schau dir SOWOHL den Spaltennamen ALS AUCH die Beispielwerte an. Die Werte sind wichtiger.
2. Wenn eine Spalte sowohl Vor- als auch Nachnamen enthält (z.B. "Mag. Stefan Holzer"), nutze full_name.
3. has_no_header_row = true wenn die "Spaltenheader" wie echte Daten aussehen (Namen, Firmen, E-Mails).
4. _ignore verwenden für: interne IDs, Zeilennummern, leere Spalten, technische Felder, Land-Kürzel als Prefix.
5. custom_field_suggestions NUR für Daten die für einen Trainer/Coach wertvoll sind:
   - Sinnvoll: Interessen, Beruf, Region/Ort, Spezialisierung, Sprache, Notizen-Kategorie
   - NICHT sinnvoll: interne Nummern, Zeilen-IDs, Timestamps, technische Flags, Spalten mit < 10% Befüllung
6. Antworte ausschließlich mit validem JSON, kein Markdown, keine Erklärungen außerhalb des JSON.

Ausgabeformat (strikt):
{
  "has_no_header_row": false,
  "language_detected": "de",
  "mappings": { "0": "organization_name", "1": "full_name", "2": "_ignore" },
  "custom_field_suggestions": [
    { "column_index": 5, "column_header": "Region", "suggested_name": "Region", "type": "text",
      "reason": "Enthält Bundesland-Kürzel – nützlich für regionale Segmentierung" }
  ],
  "notes": "Optionale Anmerkung zur Datei"
}
""".strip()


@router.post("/import/analyze-columns")
async def analyze_columns(payload: AnalyzeColumnsRequest, request: Request):
    """
    GPT-4o analysiert Spalten anhand von Namen UND Beispielwerten.
    Erkennt fehlende Header-Zeilen und schlägt sinnvolle Custom Fields vor.
    """
    if not payload.headers:
        raise HTTPException(400, "headers darf nicht leer sein")

    # Kompakte Tabelle für den Prompt aufbauen
    table_lines = []
    for i, header in enumerate(payload.headers):
        samples = [row[i] for row in payload.sample_rows if i < len(row) and row[i].strip()][:5]
        samples_str = " | ".join(samples) if samples else "(leer)"
        table_lines.append(f"  Spalte {i}: Header=«{header}»  Beispiele: {samples_str}")

    existing_note = ""
    if payload.existing_custom_field_names:
        existing_note = f"\nBereits vorhandene eigene Felder (nicht nochmals vorschlagen): {', '.join(payload.existing_custom_field_names)}"

    user_msg = f"Analysiere diese {len(payload.headers)} Spalten:{existing_note}\n\n" + "\n".join(table_lines)

    try:
        response = _oai.chat.completions.create(
            model=_ANALYZE_MODEL,
            messages=[
                {"role": "system", "content": _ANALYZE_SYSTEM_PROMPT},
                {"role": "user",   "content": user_msg},
            ],
            response_format={"type": "json_object"},
            temperature=0.1,
            max_tokens=1500,
        )
        raw_json = response.choices[0].message.content or "{}"
        result = json.loads(raw_json)
        analyze_tokens_in  = response.usage.prompt_tokens     if response.usage else 0
        analyze_tokens_out = response.usage.completion_tokens if response.usage else 0
    except json.JSONDecodeError as e:
        log.error("analyze-columns: Ungültiges JSON von GPT: %s", e)
        raise HTTPException(500, "GPT hat kein valides JSON zurückgegeben")
    except Exception as e:
        log.error("analyze-columns: GPT-Fehler: %s", e, exc_info=True)
        raise HTTPException(500, f"KI-Analyse fehlgeschlagen: {str(e)}")

    # Mapping: Keys sicherstellen dass sie Integers sind (GPT gibt manchmal Strings zurück)
    raw_mappings = result.get("mappings", {})
    mappings = {str(k): str(v) for k, v in raw_mappings.items()}

    return {
        "has_no_header_row":        bool(result.get("has_no_header_row", False)),
        "language_detected":        result.get("language_detected", "unknown"),
        "mappings":                 mappings,
        "custom_field_suggestions": result.get("custom_field_suggestions", []),
        "notes":                    result.get("notes", ""),
        "model_used":               _ANALYZE_MODEL,
        "tokens_input":             analyze_tokens_in,
        "tokens_output":            analyze_tokens_out,
    }


_EXTRACT_SYSTEM_PROMPT = """
Du extrahierst Kontaktdaten aus beliebigen Dokumenten für einen Trainer/Coach.

Regeln:
1. Suche nach PERSONEN und ORGANISATIONEN — ignoriere KPIs, Summen, Metadaten, Diagramme, leere Zeilen, interne IDs, Zeilennummern.
2. Erkenne Beziehungen: welche Person gehört zu welcher Organisation?
3. SCHLIESSE AUS: Personen/Firmen die in exclude_identity stehen (das ist der Trainer selbst).
4. Bei Rechnungen: die Absender-Seite ist meist der Trainer (ausschließen), die Empfänger-Seite ist der Kunde (extrahieren).
5. Wenn ein Name Titel enthält (Mag., Dr., DI, Prof.) → in full_name übernehmen.
6. Gib nur echte Kontaktdaten zurück — keine Platzhalter, keine Beispieldaten.
7. confidence: 0.0-1.0 (wie sicher bist du dass das ein echter Kontakt ist)

CUSTOM FIELDS — WICHTIG:
8. Die Standardfelder sind: first_name, last_name, email, phone, organization_name, organization_role, date_of_birth, notes.
9. ALLE weiteren Informationen pro Kontakt gehören in "custom_fields" — z.B. Adresse, Stadt, PLZ, Land, Region, Branche, Beruf, Titel, Website, Kundennummer, Abteilung, Geschlecht, Sprache, etc.
10. Wenn der Trainer bereits eigene Custom-Felder hat (werden im User-Prompt angegeben), verwende EXAKT deren Namen als Keys.
11. Für neue Felder: verwende aussagekräftige, kurze deutsche Feldnamen als Keys (z.B. "Straße", "PLZ", "Stadt", "Land", "Branche", "Website").
12. Adressfelder IMMER einzeln aufteilen: "Straße", "PLZ", "Stadt", "Land" — NICHT als ein zusammengesetztes Feld.

Antworte ausschließlich mit validem JSON:
{
  "contacts": [
    {
      "first_name": "Stefan",
      "last_name": "Holzer",
      "full_name": null,
      "email": "stefan@example.at",
      "phone": "+43 664 123 45 67",
      "organization_name": "Ärztekammer Vorarlberg",
      "organization_role": "Vorstandsmitglied",
      "date_of_birth": null,
      "notes": null,
      "custom_fields": {
        "Straße": "Schulgasse 17",
        "PLZ": "6850",
        "Stadt": "Dornbirn",
        "Land": "Österreich",
        "Branche": "Gesundheitswesen"
      },
      "confidence": 0.95,
      "source_hint": "Zeile 5"
    }
  ],
  "organizations": [
    {
      "name": "Ärztekammer Vorarlberg",
      "website": null,
      "address": "Schulgasse 17, 6850 Dornbirn"
    }
  ],
  "chunk_notes": "200 Zeilen verarbeitet, 45 Kontakte gefunden, 12 Zeilen waren KPI-Daten"
}
""".strip()


def _rows_to_text_table(rows: list[list]) -> str:
    """Konvertiert Zeilen in lesbaren Text für GPT."""
    lines = []
    for i, row in enumerate(rows):
        cells = [str(c).strip() if c is not None else "" for c in row]
        # Leere Zeilen überspringen
        if not any(cells):
            continue
        lines.append(f"Zeile {i+1}: " + " | ".join(cells))
    return "\n".join(lines)


def _extract_pdf_text(content_bytes: bytes) -> str:
    """Extrahiert Text aus PDF via pdfplumber."""
    try:
        import pdfplumber
        import io as _io
        with pdfplumber.open(_io.BytesIO(content_bytes)) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n--- Seite ---\n\n".join(pages)
    except Exception as e:
        log.warning("PDF-Extraktion fehlgeschlagen: %s", e)
        return ""


def _extract_docx_text(content_bytes: bytes) -> str:
    """Extrahiert Text + Tabellen aus Word-Dokumenten (.docx)."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content_bytes))
        parts = []
        # Absätze
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        # Tabellen
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        log.warning("DOCX-Extraktion fehlgeschlagen: %s", e)
        return ""


def _extract_rtf_text(content_bytes: bytes) -> str:
    """Extrahiert Text aus RTF-Dateien."""
    try:
        from striprtf.striprtf import rtf_to_text
        rtf_content = content_bytes.decode("utf-8", errors="replace")
        return rtf_to_text(rtf_content)
    except Exception as e:
        log.warning("RTF-Extraktion fehlgeschlagen: %s", e)
        return ""


def _extract_odt_text(content_bytes: bytes) -> str:
    """Extrahiert Text aus OpenDocument-Textdateien (.odt)."""
    try:
        from odf.opendocument import load as odf_load
        from odf.text import P
        from odf import teletype
        doc = odf_load(io.BytesIO(content_bytes))
        parts = []
        for p in doc.getElementsByType(P):
            text = teletype.extractText(p).strip()
            if text:
                parts.append(text)
        return "\n".join(parts)
    except Exception as e:
        log.warning("ODT-Extraktion fehlgeschlagen: %s", e)
        return ""


def _extract_eml_text(content_bytes: bytes) -> str:
    """Extrahiert Absender, Empfänger und Text aus E-Mail-Dateien (.eml)."""
    try:
        import email
        from email import policy
        msg = email.message_from_bytes(content_bytes, policy=policy.default)
        parts = []
        # Header-Infos (Kontaktdaten!)
        for header in ["From", "To", "Cc", "Bcc", "Reply-To"]:
            val = msg.get(header)
            if val:
                parts.append(f"{header}: {val}")
        # Subject
        if msg.get("Subject"):
            parts.append(f"Betreff: {msg['Subject']}")
        # Body
        body = msg.get_body(preferencelist=("plain", "html"))
        if body:
            text = body.get_content()
            if isinstance(text, bytes):
                text = text.decode("utf-8", errors="replace")
            parts.append(f"\n--- E-Mail-Text ---\n{text}")
        return "\n".join(parts)
    except Exception as e:
        log.warning("EML-Extraktion fehlgeschlagen: %s", e)
        return ""


def _parse_vcf(content_bytes: bytes) -> list[dict]:
    """Parst vCard-Dateien mechanisch — kein GPT nötig.
    Gibt eine Liste von Kontakt-Dicts zurück (gleiches Schema wie GPT-Extraktion)."""
    try:
        import vobject
        text = content_bytes.decode("utf-8", errors="replace")
        contacts = []
        for vcard in vobject.readComponents(text):
            contact: dict = {}
            # Name
            if hasattr(vcard, "n"):
                n = vcard.n.value
                contact["first_name"] = n.given or ""
                contact["last_name"] = n.family or ""
            elif hasattr(vcard, "fn"):
                parts = vcard.fn.value.split(" ", 1)
                contact["first_name"] = parts[0] if parts else ""
                contact["last_name"] = parts[1] if len(parts) > 1 else ""
            # E-Mail
            if hasattr(vcard, "email"):
                contact["email"] = vcard.email.value
            # Telefon
            if hasattr(vcard, "tel"):
                contact["phone"] = vcard.tel.value
            # Organisation
            if hasattr(vcard, "org"):
                org_val = vcard.org.value
                contact["organization"] = org_val[0] if isinstance(org_val, list) else str(org_val)
            # Titel/Position
            if hasattr(vcard, "title"):
                contact["position"] = vcard.title.value
            # Adresse
            if hasattr(vcard, "adr"):
                adr = vcard.adr.value
                if adr.street:
                    contact["street"] = adr.street
                if adr.city:
                    contact["city"] = adr.city
                if adr.code:
                    contact["zip"] = adr.code
                if adr.country:
                    contact["country"] = adr.country
            # Website
            if hasattr(vcard, "url"):
                contact["website"] = vcard.url.value
            # Nur hinzufügen wenn mindestens Name oder E-Mail vorhanden
            if contact.get("first_name") or contact.get("last_name") or contact.get("email"):
                contacts.append(contact)
        log.info("vCard: %d Kontakte mechanisch geparst", len(contacts))
        return contacts
    except Exception as e:
        log.warning("vCard-Parsing fehlgeschlagen: %s", e)
        return []


def _pdf_pages_to_images(content_bytes: bytes, dpi: int = 150) -> list[str]:
    """Rendert PDF-Seiten als PNG-Bilder und gibt Base64-Strings zurück."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=content_bytes, filetype="pdf")
        images = []
        for page in doc:
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            images.append(base64.b64encode(img_bytes).decode("utf-8"))
        doc.close()
        log.info("PDF→Bilder: %d Seiten gerendert (%d DPI)", len(images), dpi)
        return images
    except ImportError:
        log.error("PyMuPDF nicht installiert — PDF-Vision-Fallback nicht verfügbar")
        return []
    except Exception as e:
        log.error("PDF→Bilder fehlgeschlagen: %s", e)
        return []


async def _call_gpt_extract(content_text: str | None, image_b64: str | None,
                             image_mime: str | None, exclude_identity: dict,
                             available_fields: list[str],
                             model: str = "gpt-4.1",
                             _max_tokens: int = 16000,
                             _attempt: int = 1,
                             _prev_tokens_in: int = 0,
                             _prev_tokens_out: int = 0) -> dict:
    """Ruft GPT zur Kontakt-Extraktion auf. Entweder Text oder Bild.
    Bei Truncation (finish_reason=='length') wird automatisch mit höherem
    Token-Limit erneut versucht (max 2 Versuche)."""

    # Exclude-Hinweis aufbauen
    exclude_parts = []
    if exclude_identity.get("names"):
        exclude_parts.append("Namen: " + ", ".join(exclude_identity["names"]))
    if exclude_identity.get("emails"):
        exclude_parts.append("E-Mails: " + ", ".join(exclude_identity["emails"]))
    if exclude_identity.get("company_name"):
        exclude_parts.append("Firma: " + exclude_identity["company_name"])
    exclude_note = ("AUSSCHLIESSEN (das ist der Trainer selbst): " + " | ".join(exclude_parts)) if exclude_parts else ""

    # System-Felder rausfiltern, nur echte Custom Fields des Trainers anzeigen
    system_fields = {"first_name", "last_name", "full_name", "email", "phone",
                     "date_of_birth", "organization_name", "organization_role",
                     "status", "notes", "doi_confirmed_at", "_ignore"}
    custom_only = [f for f in available_fields if f not in system_fields]
    if custom_only:
        fields_note = "Bereits vorhandene Custom-Felder des Trainers (verwende EXAKT diese Namen wenn passend): " + ", ".join(custom_only)
    else:
        fields_note = ""

    user_content: list = []

    if image_b64 and image_mime:
        user_content.append({
            "type": "text",
            "text": f"Extrahiere alle Kontakte aus diesem Dokument/Bild.\n{exclude_note}\n{fields_note}"
        })
        user_content.append({
            "type": "image_url",
            "image_url": {"url": f"data:{image_mime};base64,{image_b64}", "detail": "high"}
        })
    else:
        user_content.append({
            "type": "text",
            "text": f"{exclude_note}\n{fields_note}\n\nInhalt:\n{content_text}"
        })

    response = _oai.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": _EXTRACT_SYSTEM_PROMPT},
            {"role": "user",   "content": user_content},
        ],
        response_format={"type": "json_object"},
        temperature=0.1,
        max_tokens=_max_tokens,
    )

    finish_reason = response.choices[0].finish_reason
    raw = response.choices[0].message.content or "{}"

    # Echte Token-Zahlen aus der API-Response
    this_tokens_in  = response.usage.prompt_tokens     if response.usage else 0
    this_tokens_out = response.usage.completion_tokens if response.usage else 0

    # Truncation-Erkennung: GPT wurde mitten im JSON abgeschnitten
    if finish_reason == "length":
        log.warning("GPT-Extraktion truncated (finish_reason=length, max_tokens=%d, attempt=%d)",
                     _max_tokens, _attempt)
        # Einmal mit doppeltem Limit retry (max 2 Versuche)
        if _attempt < 2:
            return await _call_gpt_extract(
                content_text=content_text, image_b64=image_b64,
                image_mime=image_mime, exclude_identity=exclude_identity,
                available_fields=available_fields,
                model=model,
                _max_tokens=min(_max_tokens * 2, 16384),
                _attempt=_attempt + 1,
                _prev_tokens_in=_prev_tokens_in + this_tokens_in,
                _prev_tokens_out=_prev_tokens_out + this_tokens_out,
            )
        # Nach 2 Versuchen: so viel wie möglich retten
        log.error("GPT-Extraktion nach %d Versuchen immer noch truncated — versuche partielle Rettung", _attempt)
        try:
            data = json.loads(raw)
        except json.JSONDecodeError:
            data = {"contacts": [], "organizations": [],
                    "chunk_notes": f"Antwort wurde nach {_max_tokens} Tokens abgeschnitten und konnte nicht geparst werden."}
        data["_tokens_in"]  = _prev_tokens_in  + this_tokens_in
        data["_tokens_out"] = _prev_tokens_out + this_tokens_out
        return data

    data = json.loads(raw)
    data["_tokens_in"]  = _prev_tokens_in  + this_tokens_in
    data["_tokens_out"] = _prev_tokens_out + this_tokens_out
    return data


# Semaphore: max 5 parallele GPT-Calls (Rate-Limit-Schutz)
_gpt_semaphore = asyncio.Semaphore(5)


def _log_gpt_result(result: dict, source: str, file_name: str) -> None:
    """Loggt GPT-Ergebnis direkt nach dem Call — bevor irgendetwas gefiltert wird."""
    contacts = result.get("contacts", [])
    orgs = result.get("organizations", [])
    notes = result.get("chunk_notes", "")
    if contacts:
        log.info("extract-contacts [%s] %s: %d Kontakte, %d Orgs gefunden. Notes: %s",
                 source, file_name, len(contacts), len(orgs), notes or "(keine)")
    else:
        log.warning("extract-contacts [%s] %s: 0 Kontakte gefunden. Notes: %s",
                    source, file_name, notes or "(keine)")


async def _extract_chunk(content_text: str | None, image_b64: str | None,
                          image_mime: str | None, exclude_identity: dict,
                          available_fields: list[str], label: str,
                          file_name: str, model: str = "gpt-4.1") -> dict:
    """Einen Chunk mit Semaphore-Schutz extrahieren."""
    async with _gpt_semaphore:
        result = await _call_gpt_extract(
            content_text=content_text,
            image_b64=image_b64, image_mime=image_mime,
            exclude_identity=exclude_identity,
            available_fields=available_fields,
            model=model,
        )
        _log_gpt_result(result, label, file_name)
        return result


def _collect_results(results: list[dict]) -> tuple[list[dict], list[dict], list[str], int, int]:
    """Sammelt Kontakte, Orgs, Notes und Token-Summen aus mehreren Chunk-Ergebnissen."""
    contacts, orgs, notes = [], [], []
    tokens_in, tokens_out = 0, 0
    for r in results:
        contacts.extend(r.get("contacts", []))
        orgs.extend(r.get("organizations", []))
        if r.get("chunk_notes"):
            notes.append(r["chunk_notes"])
        tokens_in  += r.get("_tokens_in",  0)
        tokens_out += r.get("_tokens_out", 0)
    return contacts, orgs, notes, tokens_in, tokens_out


async def _process_sheet_parallel(
    sheet: dict,
    exclude_identity: dict,
    available_fields: list,
    file_name: str,
    model: str,
) -> tuple[list, list, list, int, int]:
    """Verarbeitet ein einzelnes Sheet (alle Chunks parallel). Wird von sync + async Endpoint genutzt."""
    rows = sheet.get("rows", [])
    headers = sheet.get("headers", [])
    sheet_name = sheet.get("name", "")

    all_rows = [headers] + rows if headers else rows
    chunk_size = 50
    total_chunks = (len(all_rows) + chunk_size - 1) // chunk_size

    tasks = []
    for chunk_start in range(0, len(all_rows), chunk_size):
        chunk_idx = chunk_start // chunk_size + 1
        chunk = all_rows[chunk_start:chunk_start + chunk_size]
        text = f"Sheet: {sheet_name}\n" + _rows_to_text_table(chunk)
        tasks.append(_extract_chunk(
            content_text=text, image_b64=None, image_mime=None,
            exclude_identity=exclude_identity,
            available_fields=available_fields,
            label=f"{sheet_name}-chunk-{chunk_idx}/{total_chunks}",
            file_name=file_name,
            model=model,
        ))

    results = await asyncio.gather(*tasks)
    c, o, n, ti, to = _collect_results(list(results))
    # Notes mit Sheet-Prefix
    n = [f"[{sheet_name}] {note}" for note in n]
    return c, o, n, ti, to


async def _run_extraction(content_bytes: bytes, payload: ExtractContactsRequest) -> dict:
    """
    Kernlogik der Kontaktextraktion — verwendbar von sync UND async Endpoint.
    Gibt das fertige Result-Dict zurück (contacts, organizations, notes, tokens, ...).
    Sheets werden PARALLEL verarbeitet (max 5 concurrent GPT-Calls via Semaphore).
    """
    ft = payload.file_type.lower().lstrip(".")

    all_contacts: list[dict] = []
    all_orgs: list[dict] = []
    all_notes: list[str] = []
    all_tokens_in:  int = 0
    all_tokens_out: int = 0
    vision_tokens_in:  int = 0
    vision_tokens_out: int = 0

    # ── Bilder → GPT Vision ────────────────────────────────────────────
    IMAGE_TYPES = {"jpg": "image/jpeg", "jpeg": "image/jpeg",
                   "png": "image/png", "webp": "image/webp", "heic": "image/heic",
                   "tif": "image/tiff", "tiff": "image/tiff", "bmp": "image/bmp"}
    if ft in IMAGE_TYPES:
        result = await _extract_chunk(
            content_text=None,
            image_b64=payload.file_content_base64, image_mime=IMAGE_TYPES[ft],
            exclude_identity=payload.exclude_identity,
            available_fields=payload.available_fields,
            label="vision", file_name=payload.file_name,
            model=payload.vision_model,
        )
        all_contacts.extend(result.get("contacts", []))
        all_orgs.extend(result.get("organizations", []))
        if result.get("chunk_notes"):
            all_notes.append(result["chunk_notes"])
        vision_tokens_in  += result.get("_tokens_in",  0)
        vision_tokens_out += result.get("_tokens_out", 0)

    # ── PDF → Text → chunked parallel → Vision-Fallback ────────────────
    elif ft == "pdf":
        pdf_text = _extract_pdf_text(content_bytes)
        if pdf_text.strip():
            log.info("extract-contacts [pdf] %s: %d Zeichen Text extrahiert.", payload.file_name, len(pdf_text))
            chunk_size = 6000
            text_chunks = [pdf_text[i:i+chunk_size] for i in range(0, len(pdf_text), chunk_size)]
            total = len(text_chunks)
            tasks = [
                _extract_chunk(
                    content_text=chunk, image_b64=None, image_mime=None,
                    exclude_identity=payload.exclude_identity,
                    available_fields=payload.available_fields,
                    label=f"pdf-chunk-{i+1}/{total}", file_name=payload.file_name,
                    model=payload.model,
                )
                for i, chunk in enumerate(text_chunks)
            ]
            results = await asyncio.gather(*tasks)
            c, o, n, ti, to = _collect_results(list(results))
            all_contacts.extend(c); all_orgs.extend(o); all_notes.extend(n)
            all_tokens_in += ti; all_tokens_out += to

        if not all_contacts:
            log.info("extract-contacts [pdf-vision-fallback] %s: Text lieferte 0 Kontakte, versuche Vision…", payload.file_name)
            page_images = _pdf_pages_to_images(content_bytes)
            if page_images:
                vision_tasks = [
                    _extract_chunk(
                        content_text=None, image_b64=img_b64, image_mime="image/png",
                        exclude_identity=payload.exclude_identity,
                        available_fields=payload.available_fields,
                        label=f"pdf-vision-{i+1}/{len(page_images)}",
                        file_name=payload.file_name, model=payload.vision_model,
                    )
                    for i, img_b64 in enumerate(page_images)
                ]
                vision_results = await asyncio.gather(*vision_tasks)
                vc, vo, vn, vti, vto = _collect_results(list(vision_results))
                vision_tokens_in += vti; vision_tokens_out += vto
                if vc:
                    all_contacts = vc; all_orgs = vo
                    all_notes = [f"[Vision-Fallback] {n}" for n in vn]

    # ── vCard → mechanisches Parsing (kein GPT nötig) ──────────────────
    elif ft == "vcf":
        vcf_contacts = _parse_vcf(content_bytes)
        all_contacts.extend(vcf_contacts)
        if vcf_contacts:
            seen_orgs = set()
            for c in vcf_contacts:
                org_name = c.pop("organization", None)
                if org_name and org_name not in seen_orgs:
                    seen_orgs.add(org_name)
                    all_orgs.append({"name": org_name})
            all_notes.append(f"{len(vcf_contacts)} Kontakte aus vCard importiert.")
        else:
            all_notes.append("vCard-Datei enthielt keine auswertbaren Kontakte.")

    # ── Dokumente (DOCX, RTF, ODT, EML) → Text → GPT ─────────────────
    elif ft in ("docx", "rtf", "odt", "eml"):
        extractors = {"docx": _extract_docx_text, "rtf": _extract_rtf_text,
                      "odt": _extract_odt_text, "eml": _extract_eml_text}
        doc_text = extractors[ft](content_bytes)
        if doc_text.strip():
            chunk_size = 6000
            text_chunks = [doc_text[i:i+chunk_size] for i in range(0, len(doc_text), chunk_size)]
            total = len(text_chunks)
            tasks = [
                _extract_chunk(
                    content_text=chunk, image_b64=None, image_mime=None,
                    exclude_identity=payload.exclude_identity,
                    available_fields=payload.available_fields,
                    label=f"{ft}-chunk-{i+1}/{total}", file_name=payload.file_name,
                    model=payload.model,
                )
                for i, chunk in enumerate(text_chunks)
            ]
            results = await asyncio.gather(*tasks)
            c, o, n, ti, to = _collect_results(list(results))
            all_contacts.extend(c); all_orgs.extend(o); all_notes.extend(n)
            all_tokens_in += ti; all_tokens_out += to
        else:
            all_notes.append(f"{ft.upper()}-Datei enthielt keinen lesbaren Text.")

    # ── CSV / TXT / TSV → Zeilen → chunked parallel ────────────────────
    elif ft in ("csv", "txt", "tsv"):
        sheet = parse_csv_bytes(content_bytes)
        sheet_results = await asyncio.gather(_process_sheet_parallel(
            sheet, payload.exclude_identity, payload.available_fields, payload.file_name, payload.model
        ))
        c, o, n, ti, to = sheet_results[0]
        all_contacts.extend(c); all_orgs.extend(o); all_notes.extend(n)
        all_tokens_in += ti; all_tokens_out += to

    # ── Excel / ODS → alle Sheets PARALLEL ────────────────────────────
    elif ft in ("xlsx", "xls", "ods", "xlsm"):
        sheets = parse_excel_bytes(content_bytes)
        log.info("extract-contacts [xlsx] %s: %d Sheets werden parallel verarbeitet", payload.file_name, len(sheets))

        # Alle Sheets gleichzeitig starten (Semaphore begrenzt concurrent GPT-Calls)
        sheet_tasks = [
            _process_sheet_parallel(
                sheet, payload.exclude_identity, payload.available_fields, payload.file_name, payload.model
            )
            for sheet in sheets
        ]
        sheet_results = await asyncio.gather(*sheet_tasks)

        for c, o, n, ti, to in sheet_results:
            all_contacts.extend(c); all_orgs.extend(o); all_notes.extend(n)
            all_tokens_in += ti; all_tokens_out += to

    else:
        raise HTTPException(415, f"Nicht unterstütztes Format: {ft}")

    log.info("extract-contacts FERTIG [%s] %s: %d Kontakte, %d Orgs, text=%d/%d tokens, vision=%d/%d tokens",
             ft, payload.file_name, len(all_contacts), len(all_orgs),
             all_tokens_in, all_tokens_out, vision_tokens_in, vision_tokens_out)

    return {
        "contacts":            all_contacts,
        "organizations":       all_orgs,
        "notes":               " | ".join(all_notes),
        "total_found":         len(all_contacts),
        "tokens_input":        all_tokens_in,
        "tokens_output":       all_tokens_out,
        "vision_tokens_input":  vision_tokens_in,
        "vision_tokens_output": vision_tokens_out,
        "vision_used":         vision_tokens_in > 0,
    }


@router.post("/import/extract-contacts")
async def extract_contacts(payload: ExtractContactsRequest, request: Request):
    """
    Synchroner Endpoint — wartet auf Ergebnis (für kleine Dateien / Rückwärtskompatibilität).
    Nutzt intern _run_extraction() mit parallelen Sheets.
    Für große Multi-Sheet Excel: /import/extract-contacts-async verwenden.
    """
    content_bytes = base64.b64decode(payload.file_content_base64)
    try:
        return await _run_extraction(content_bytes, payload)
    except HTTPException:
        raise
    except Exception as e:
        log.error("extract-contacts: Fehler: %s", e, exc_info=True)
        raise HTTPException(500, f"Extraktion fehlgeschlagen: {str(e)}")


async def _extract_contacts_background(task_id: str, payload: ExtractContactsAsyncRequest):
    """
    Hintergrund-Coroutine für async Extraktion.
    Läuft unabhängig vom HTTP-Request — kein PHP-Timeout-Problem.
    Speichert Ergebnis in _import_tasks und ruft optional Callback auf.
    """
    try:
        content_bytes = base64.b64decode(payload.file_content_base64)
        result = await _run_extraction(content_bytes, payload)

        _import_tasks[task_id]["status"] = "complete"
        _import_tasks[task_id]["result"] = result
        log.info("import-async [%s] job=%d: Fertig — %d Kontakte",
                 task_id[:8], payload.job_id, result.get("total_found", 0))

        # Callback an Laravel
        callback_url = payload.callback_url
        if not callback_url and payload.laravel_base_url:
            callback_url = payload.laravel_base_url.rstrip("/") + "/api/ai/import/complete"

        if callback_url:
            try:
                async with httpx.AsyncClient(timeout=30) as client:
                    resp = await client.post(
                        callback_url,
                        json={"job_id": payload.job_id, "task_id": task_id, **result},
                        headers={"Authorization": f"Bearer {_GATEWAY_SECRET}"},
                    )
                    if resp.status_code == 200:
                        log.info("import-async [%s] job=%d: Callback OK", task_id[:8], payload.job_id)
                    else:
                        log.warning("import-async [%s] job=%d: Callback HTTP %d — %s",
                                    task_id[:8], payload.job_id, resp.status_code, resp.text[:200])
            except Exception as cb_err:
                log.warning("import-async [%s] job=%d: Callback fehlgeschlagen: %s — Ergebnis via Polling abrufbar",
                            task_id[:8], payload.job_id, cb_err)

    except Exception as e:
        log.error("import-async [%s] job=%d: Extraktion fehlgeschlagen: %s", task_id[:8], payload.job_id, e, exc_info=True)
        _import_tasks[task_id]["status"] = "failed"
        _import_tasks[task_id]["error"] = str(e)

        # Failure-Callback
        fail_url = payload.callback_url
        if not fail_url and payload.laravel_base_url:
            fail_url = payload.laravel_base_url.rstrip("/") + "/api/ai/import/failed"
        elif fail_url:
            fail_url = fail_url.replace("/import/complete", "/import/failed")

        if fail_url:
            try:
                async with httpx.AsyncClient(timeout=15) as client:
                    await client.post(
                        fail_url,
                        json={"job_id": payload.job_id, "task_id": task_id, "reason": str(e)},
                        headers={"Authorization": f"Bearer {_GATEWAY_SECRET}"},
                    )
            except Exception:
                pass


@router.post("/import/extract-contacts-async")
async def extract_contacts_async(payload: ExtractContactsAsyncRequest, request: Request):
    """
    Async Endpoint — kehrt sofort zurück (< 1s), verarbeitet im Hintergrund.
    Gateway läuft auf VPS ohne PHP-Timeout-Einschränkung.
    Ergebnis via Callback (production) oder GET /import/status/{task_id} (dev/fallback).
    """
    task_id = str(uuid.uuid4())
    _import_tasks[task_id] = {
        "status": "processing",
        "result": None,
        "error":  None,
        "job_id": payload.job_id,
        "file_name": payload.file_name,
        "created_at": time.time(),
    }
    log.info("import-async [%s] job=%d: Starte Hintergrundverarbeitung für '%s'",
             task_id[:8], payload.job_id, payload.file_name)

    # Fire-and-forget als asyncio Task (läuft im selben Event-Loop)
    asyncio.create_task(_extract_contacts_background(task_id, payload))

    return {"status": "accepted", "task_id": task_id, "job_id": payload.job_id}


@router.get("/import/status/{task_id}")
async def get_import_status(task_id: str, request: Request):
    """
    Polling-Fallback für lokale Dev-Umgebungen wo Callbacks nicht funktionieren.
    Gibt Extraktion-Status + Ergebnis zurück sobald fertig.
    Cleanup: Tasks älter als 2h werden entfernt.
    """
    # Cleanup alter Tasks (>2h)
    now = time.time()
    expired = [tid for tid, t in _import_tasks.items() if now - t.get("created_at", now) > 7200]
    for tid in expired:
        _import_tasks.pop(tid, None)

    task = _import_tasks.get(task_id)
    if not task:
        raise HTTPException(404, f"Task '{task_id}' nicht gefunden (abgelaufen oder ungültig)")

    return {
        "task_id":   task_id,
        "status":    task["status"],           # "processing" | "complete" | "failed"
        "job_id":    task.get("job_id"),
        "file_name": task.get("file_name"),
        "result":    task.get("result"),       # None wenn noch in Bearbeitung
        "error":     task.get("error"),
    }


def _normalize_single(raw: str, fallback: Optional[str], candidates: list[str], phonenumbers) -> dict:
    """Normalisiert eine einzelne Telefonnummer."""
    if not raw or not raw.strip():
        return {"e164": None, "country": None, "confidence": "none", "possible_countries": [], "national_format": None}

    raw = raw.strip()

    # Schritt 1: Direkt parsen (falls +XX vorhanden)
    try:
        parsed = phonenumbers.parse(raw, None)
        if phonenumbers.is_valid_number(parsed):
            country = phonenumbers.region_code_for_number(parsed)
            return {
                "e164": phonenumbers.format_number(parsed, phonenumbers.PhoneNumberFormat.E164),
                "country": country,
                "confidence": "high",
                "possible_countries": [country],
                "national_format": phonenumbers.format_number(parsed, phonenumbers.PhoneNumberFormat.NATIONAL),
            }
    except Exception:
        pass

    # Schritt 2: Fallback-Land direkt versuchen
    if fallback:
        try:
            parsed = phonenumbers.parse(raw, fallback)
            if phonenumbers.is_valid_number(parsed):
                country = phonenumbers.region_code_for_number(parsed)
                return {
                    "e164": phonenumbers.format_number(parsed, phonenumbers.PhoneNumberFormat.E164),
                    "country": country,
                    "confidence": "medium",
                    "possible_countries": [country],
                    "national_format": phonenumbers.format_number(parsed, phonenumbers.PhoneNumberFormat.NATIONAL),
                }
        except Exception:
            pass

    # Schritt 2b: Führende 0 wieder hinzufügen (Excel schneidet sie ab)
    # z.B. "677 125 511" → "0677 125 511" → valide AT-Mobilnummer
    digits_only = re.sub(r'[^\d]', '', raw)
    if fallback and digits_only and not digits_only.startswith('0') and not raw.startswith('+'):
        raw_with_zero = "0" + raw.lstrip()
        try:
            parsed = phonenumbers.parse(raw_with_zero, fallback)
            if phonenumbers.is_valid_number(parsed):
                country = phonenumbers.region_code_for_number(parsed)
                log.debug("Telefon: führende 0 ergänzt: '%s' → '%s' (%s)", raw, raw_with_zero, country)
                return {
                    "e164": phonenumbers.format_number(parsed, phonenumbers.PhoneNumberFormat.E164),
                    "country": country,
                    "confidence": "medium",
                    "possible_countries": [country],
                    "national_format": phonenumbers.format_number(parsed, phonenumbers.PhoneNumberFormat.NATIONAL),
                }
        except Exception:
            pass

    # Schritt 3: Alle Kandidaten-Länder durchprobieren
    possible = []
    for cc in candidates:
        try:
            parsed = phonenumbers.parse(raw, cc)
            if phonenumbers.is_valid_number(parsed):
                e164 = phonenumbers.format_number(parsed, phonenumbers.PhoneNumberFormat.E164)
                possible.append({"country": cc, "e164": e164, "parsed": parsed})
        except Exception:
            continue

    if not possible:
        return {"e164": None, "country": None, "confidence": "low", "possible_countries": [], "national_format": raw}

    if len(possible) == 1:
        p = possible[0]
        return {
            "e164": p["e164"],
            "country": p["country"],
            "confidence": "high",
            "possible_countries": [p["country"]],
            "national_format": phonenumbers.format_number(p["parsed"], phonenumbers.PhoneNumberFormat.NATIONAL),
        }

    confidence = "medium" if len(possible) <= 3 else "low"
    # Bestes Ergebnis: Fallback-Land bevorzugen, sonst erstes
    best = next((p for p in possible if p["country"] == fallback), possible[0])
    return {
        "e164": best["e164"],
        "country": best["country"],
        "confidence": confidence,
        "possible_countries": [p["country"] for p in possible],
        "national_format": phonenumbers.format_number(best["parsed"], phonenumbers.PhoneNumberFormat.NATIONAL),
    }

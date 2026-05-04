"""
Deterministischer Applier fuer einen validierten Edit-Plan.

Eingabe : DOCX-Bytes + EditPlan (siehe edit_plan.py)
Ausgabe : DOCX-Bytes (mutierte Kopie) + Applier-Report (was angewendet wurde,
          welche Operationen geskippt wurden und warum)

Wichtige Designentscheidungen:
  * Format-Erhalt bei Paragraphs: wir veraendern den ersten Run inplace und
    leeren die uebrigen Runs — so bleiben Font, Farbe und Bold-Marker erhalten.
  * Tabellen-Rebuild: clone-deep auf der `<w:tr>`-XML-Ebene erhaelt Borders,
    Shading und Spalten-Breiten. Wir nehmen die Vorlage-Zeile mit
    row_template_index, klonen sie pro Datensatz und befuellen die Zellen.
  * Out-of-bounds-Targets werden GESKIPPT (nicht raised) — der Report listet
    sie auf, der Render geht trotzdem durch (M5: liefere mit ai_warning).
"""

from __future__ import annotations

import copy
import io
import logging
from typing import Iterable

from docx import Document
from docx.document import Document as _DocumentT
from docx.oxml.ns import qn
from docx.table import Table, _Row
from docx.text.paragraph import Paragraph

from .edit_plan import EditPlan

log = logging.getLogger("ai-gateway.invoice.applier")


def apply_edit_plan(docx_bytes: bytes, plan: EditPlan) -> tuple[bytes, dict]:
    """Wendet den Plan an und gibt (neue_bytes, report) zurueck."""
    doc: _DocumentT = Document(io.BytesIO(docx_bytes))

    paragraphs, tables = _index_body(doc)

    report = {"applied": 0, "skipped": [], "warnings": []}

    for op in plan.operations:
        try:
            if op.op == "set_text":
                _apply_set_text(paragraphs, op.target, op.value, report)
            elif op.op == "set_cell_text":
                _apply_set_cell_text(tables, op.target, op.value, report)
            elif op.op == "rebuild_table_rows":
                _apply_rebuild_table_rows(tables, op, report)
                # Indices der Tabellen koennen sich nicht aendern (wir editieren
                # innerhalb derselben <w:tbl>) — Re-Index nur fuer Paragraphs noetig,
                # die hier nicht beruehrt wurden. Kein Re-Index noetig.
            elif op.op == "delete_paragraph":
                _apply_delete_paragraph(paragraphs, op.target, report)
                # Nach Delete: die nachfolgenden Paragraphs behalten ihre alten IDs
                # in unserem Index nicht mehr — wir markieren das Slot als None.
            else:
                report["skipped"].append({"target": getattr(op, "target", "?"),
                                          "reason": f"unbekannte Operation: {op.op}"})
        except Exception as e:
            report["skipped"].append({"target": getattr(op, "target", "?"),
                                      "reason": f"{type(e).__name__}: {e}"})
            log.warning("apply_edit_plan: skip op %s: %s", op.op, e)
        else:
            report["applied"] += 1

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), report


# ── Indexer ────────────────────────────────────────────────────────────

def _index_body(doc: _DocumentT) -> tuple[list[Paragraph | None], list[Table | None]]:
    """Sammelt Paragraphs und Tabellen in der Body-Reihenfolge — gleiche Ordnung
    wie document_map.build_document_map. Listenslots werden bei Delete auf None
    gesetzt, damit nachfolgende Operationen mit den ORIGINAL-IDs noch greifen."""
    paragraphs: list[Paragraph | None] = []
    tables:     list[Table | None] = []
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            paragraphs.append(Paragraph(child, doc))
        elif child.tag == qn("w:tbl"):
            tables.append(Table(child, doc))
    return paragraphs, tables


# ── Operationen ────────────────────────────────────────────────────────

def _apply_set_text(paragraphs: list[Paragraph | None], target: str, value: str, report: dict) -> None:
    idx = _parse_p_id(target)
    if idx is None or idx >= len(paragraphs) or paragraphs[idx] is None:
        report["skipped"].append({"target": target, "reason": "Paragraph nicht gefunden"})
        return
    _replace_paragraph_text(paragraphs[idx], value)


def _apply_set_cell_text(tables: list[Table | None], target: str, value: str, report: dict) -> None:
    parsed = _parse_cell_id(target)
    if parsed is None:
        report["skipped"].append({"target": target, "reason": "Zell-ID ungueltig"})
        return
    t_idx, r_idx, c_idx = parsed
    if t_idx >= len(tables) or tables[t_idx] is None:
        report["skipped"].append({"target": target, "reason": "Tabelle nicht gefunden"})
        return
    table = tables[t_idx]
    if r_idx >= len(table.rows) or c_idx >= len(table.rows[r_idx].cells):
        report["skipped"].append({"target": target, "reason": "Zelle out-of-bounds"})
        return
    cell = table.rows[r_idx].cells[c_idx]
    # Erste Zelle, alle Paragraphs zusammenfassen — sonst entstehen Doppel-Zeilen.
    if not cell.paragraphs:
        cell.add_paragraph(value)
        return
    _replace_paragraph_text(cell.paragraphs[0], value)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def _apply_rebuild_table_rows(tables: list[Table | None], op, report: dict) -> None:
    t_idx = _parse_t_id(op.target)
    if t_idx is None or t_idx >= len(tables) or tables[t_idx] is None:
        report["skipped"].append({"target": op.target, "reason": "Tabelle nicht gefunden"})
        return
    table = tables[t_idx]
    rows = table.rows
    n_rows = len(rows)
    if op.row_template_index >= n_rows:
        report["skipped"].append({"target": op.target, "reason": "row_template_index out-of-bounds"})
        return

    template_tr = copy.deepcopy(rows[op.row_template_index]._tr)

    # Bestimme Footer-Element (Anker zum davor-Einfuegen).
    footer_anchor = None
    if op.footer_rows > 0 and (op.header_rows + op.footer_rows) <= n_rows:
        footer_anchor = rows[n_rows - op.footer_rows]._tr

    # Loesche alte Daten-Zeilen (zwischen header_rows und n_rows-footer_rows).
    data_start = op.header_rows
    data_end = n_rows - op.footer_rows
    tbl_el = table._tbl
    for tr in [r._tr for r in rows[data_start:data_end]]:
        tbl_el.remove(tr)

    # Klone Template fuer jede Datenzeile, befuelle Zellen.
    for data_row in op.data_rows:
        new_tr = copy.deepcopy(template_tr)
        if footer_anchor is not None:
            footer_anchor.addprevious(new_tr)
        else:
            tbl_el.append(new_tr)
        new_row = _Row(new_tr, table)
        for c_idx, cell_value in enumerate(data_row):
            if c_idx >= len(new_row.cells):
                break
            cell = new_row.cells[c_idx]
            if not cell.paragraphs:
                cell.add_paragraph(str(cell_value))
                continue
            _replace_paragraph_text(cell.paragraphs[0], str(cell_value))
            for extra in cell.paragraphs[1:]:
                extra._element.getparent().remove(extra._element)


def _apply_delete_paragraph(paragraphs: list[Paragraph | None], target: str, report: dict) -> None:
    idx = _parse_p_id(target)
    if idx is None or idx >= len(paragraphs) or paragraphs[idx] is None:
        report["skipped"].append({"target": target, "reason": "Paragraph nicht gefunden"})
        return
    p = paragraphs[idx]
    p._element.getparent().remove(p._element)
    paragraphs[idx] = None


# ── Helfer ─────────────────────────────────────────────────────────────

def _replace_paragraph_text(p: Paragraph, value: str) -> None:
    """Ersetzt den Text und behaelt das Format des ersten Runs."""
    runs = p.runs
    if not runs:
        p.add_run(value)
        return
    runs[0].text = value
    for run in runs[1:]:
        run.text = ""


def _parse_p_id(s: str) -> int | None:
    if not s.startswith("P"):
        return None
    try:
        return int(s[1:])
    except ValueError:
        return None


def _parse_t_id(s: str) -> int | None:
    if not s.startswith("T") or "." in s:
        return None
    try:
        return int(s[1:])
    except ValueError:
        return None


def _parse_cell_id(s: str) -> tuple[int, int, int] | None:
    if not s.startswith("T"):
        return None
    try:
        t_part, r_part, c_part = s.split(".")
        return int(t_part[1:]), int(r_part[1:]), int(c_part[1:])
    except (ValueError, IndexError):
        return None
